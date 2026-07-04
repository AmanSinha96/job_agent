"""
resume_parser.py

Phase 5.2

Resume Parser

Reads the master DOCX and
builds a structured representation
of the resume without modifying it.

Later phases will edit only the
required sections.
"""

from dataclasses import dataclass
from typing import List
from typing import Dict

from docx import Document


# ==========================================================
# Section Object
# ==========================================================

@dataclass
class ResumeSection:

    name: str

    start: int

    end: int

    paragraphs: List


# ==========================================================
# Resume Parser
# ==========================================================

class ResumeParser:

    """
    Parses DOCX resume.

    Finds:

        Summary

        Skills

        Experience

        Projects

        Certifications

        Education

    """

    # -----------------------------------------------

    def __init__(

        self,

        document: Document,

    ):

        self.document = document

        self.paragraphs = document.paragraphs

        self.sections = {}

        self.section_order = []

    # -----------------------------------------------

    @staticmethod
    def normalize(

        text,

    ):

        return (

            text

            .replace("\n", " ")

            .strip()

            .lower()

        )

    # -----------------------------------------------

    @staticmethod
    def is_heading(

        paragraph,

    ):

        text = paragraph.text.strip()

        if not text:

            return False

        style = paragraph.style.name.lower()

        if "heading" in style:

            return True

        # ALL CAPS headings

        if (

            len(text) < 60

            and text.upper() == text

        ):

            return True

        return False

    # -----------------------------------------------

    @staticmethod
    def heading_aliases():

        return {

            "summary": [

                "professional summary",

                "summary",

                "profile",

            ],

            "skills": [

                "technical skills",

                "skills",

                "core skills",

            ],

            "experience": [

                "work experience",

                "professional experience",

                "experience",

            ],

            "projects": [

                "projects",

                "ai products built",

                "personal projects",

            ],

            "education": [

                "education",

            ],

            "certifications": [

                "certifications",

                "certification",

            ],

            "achievements": [

                "achievements",

            ],

        }

    # -----------------------------------------------

    def detect_heading(

        self,

        paragraph,

    ):

        text = self.normalize(

            paragraph.text

        )

        aliases = self.heading_aliases()

        for canonical, names in aliases.items():

            for name in names:

                if text == name:

                    return canonical

        return None

    # -----------------------------------------------

    def scan_document(self):

        """
        Scan resume once.

        Identify every major section.
        """

        detected = []

        for index, paragraph in enumerate(

            self.paragraphs

        ):

            heading = self.detect_heading(

                paragraph

            )

            if heading:

                detected.append(

                    (

                        heading,

                        index,

                    )

                )

        self.section_order = detected

        return detected

    # -----------------------------------------------

    def build_sections(self):

        """
        Build section boundaries.

        Example

        SUMMARY

        paragraph 8

        ->

        ends where SKILLS begins
        """

        detected = self.scan_document()

        total = len(

            self.paragraphs

        )

        for i, (

            section_name,

            start,

        ) in enumerate(

            detected

        ):

            if i == len(detected) - 1:

                end = total

            else:

                end = detected[

                    i + 1

                ][1]

            section = ResumeSection(

                name=section_name,

                start=start,

                end=end,

                paragraphs=self.paragraphs[

                    start + 1:end

                ],

            )

            self.sections[

                section_name

            ] = section
    # -----------------------------------------------

    def get_section(

        self,

        name,

    ):

        """
        Return ResumeSection object.
        """

        if not self.sections:

            self.build_sections()

        return self.sections.get(name)

    # -----------------------------------------------

    def section_exists(

        self,

        name,

    ):

        if not self.sections:

            self.build_sections()

        return name in self.sections

    # -----------------------------------------------

    def section_text(

        self,

        name,

    ):

        """
        Returns section as plain text.
        """

        section = self.get_section(name)

        if section is None:

            return ""

        lines = []

        for paragraph in section.paragraphs:

            text = paragraph.text.strip()

            if text:

                lines.append(text)

        return "\n".join(lines)

    # -----------------------------------------------

    def section_paragraphs(

        self,

        name,

    ):

        section = self.get_section(name)

        if section is None:

            return []

        return section.paragraphs

    # -----------------------------------------------

    def get_summary(

        self,

    ):

        return self.section_text(

            "summary"

        )

    # -----------------------------------------------

    def get_skills(

        self,

    ):

        return self.section_text(

            "skills"

        )

    # -----------------------------------------------

    def get_experience(

        self,

    ):

        return self.section_text(

            "experience"

        )

    # -----------------------------------------------

    def get_projects(

        self,

    ):

        return self.section_text(

            "projects"

        )

    # -----------------------------------------------

    def get_certifications(

        self,

    ):

        return self.section_text(

            "certifications"

        )

    # -----------------------------------------------

    def get_education(

        self,

    ):

        return self.section_text(

            "education"

        )

    # -----------------------------------------------

    def get_all_sections(

        self,

    ):

        """
        Returns all parsed sections.
        """

        if not self.sections:

            self.build_sections()

        data = {}

        for section in self.sections:

            data[section] = self.section_text(

                section

            )

        return data

    # -----------------------------------------------

    def get_resume_text(

        self,

    ):

        """
        Entire resume as text.
        """

        lines = []

        for paragraph in self.paragraphs:

            text = paragraph.text.strip()

            if text:

                lines.append(text)

        return "\n".join(lines)

    # -----------------------------------------------

    def paragraph_count(

        self,

    ):

        return len(

            self.paragraphs

        )

    # -----------------------------------------------

    def section_count(

        self,

    ):

        if not self.sections:

            self.build_sections()

        return len(

            self.sections

        )

    # -----------------------------------------------

    def summary(

        self,

    ):

        """
        Human readable parser summary.
        """

        if not self.sections:

            self.build_sections()

        return {

            "paragraphs":

                self.paragraph_count(),

            "sections":

                self.section_count(),

            "detected":

                list(

                    self.sections.keys()

                )

        }

    # -----------------------------------------------

    def print_summary(

        self,

    ):

        info = self.summary()

        print()

        print("=" * 60)

        print("RESUME PARSER")

        print("=" * 60)

        print()

        print(

            "Paragraphs :",

            info["paragraphs"]

        )

        print(

            "Sections   :",

            info["sections"]

        )

        print()

        print("Detected")

        print()

        for section in info["detected"]:

            print(

                "-",

                section

            )

        print()

        print("=" * 60)

        print()
        
            # -------------------------------------------------
    # Convert skills section to list
    # -------------------------------------------------

    def get_skill_list(self):

        import re

        skills_text = self.get_skills()

        if not skills_text:
            return []

        skills = []

        for line in skills_text.split("\n"):

            line = line.strip()

            if not line:
                continue

            # remove category heading
            if ":" in line:

                _, line = line.split(
                    ":",
                    1
                )

            parts = re.split(
                r",|•|\|",
                line
            )

            for skill in parts:

                skill = skill.strip()

                if skill:

                    skills.append(skill)

        return sorted(
            list(
                set(skills)
            )
        )

    # -------------------------------------------------
    # Extract experience bullets
    # -------------------------------------------------

    def get_experience_bullets(self):

        bullets = []

        for paragraph in self.section_paragraphs(
            "experience"
        ):

            text = paragraph.text.strip()

            if not text:

                continue

            bullets.append(text)

        return bullets

    # -------------------------------------------------
    # Extract project bullets
    # -------------------------------------------------

    def get_project_bullets(self):

        bullets = []

        for paragraph in self.section_paragraphs(
            "projects"
        ):

            text = paragraph.text.strip()

            if text:

                bullets.append(text)

        return bullets

    # -------------------------------------------------
    # Find keywords already present
    # -------------------------------------------------

    def existing_keywords(self):

        import re

        text = self.get_resume_text().lower()

        words = re.findall(
            r"[a-zA-Z0-9+#.-]+",
            text
        )

        return set(words)

    # -------------------------------------------------
    # JSON export
    # -------------------------------------------------

    def to_dict(self):

        return {

            "summary":
                self.get_summary(),

            "skills":
                self.get_skills(),

            "skill_list":
                self.get_skill_list(),

            "experience":
                self.get_experience(),

            "experience_bullets":
                self.get_experience_bullets(),

            "projects":
                self.get_projects(),

            "project_bullets":
                self.get_project_bullets(),

            "certifications":
                self.get_certifications(),

            "education":
                self.get_education(),

            "resume_text":
                self.get_resume_text(),

        }

    # -------------------------------------------------
    # Pretty print
    # -------------------------------------------------

    def debug(self):

        print()

        print("=" * 70)

        print("RESUME STRUCTURE")

        print("=" * 70)

        print()

        for section in self.sections.values():

            print(

                f"{section.name.upper()}"

            )

            print(

                "-" * 40

            )

            print(

                self.section_text(

                    section.name

                )[:250]

            )

            print()

        print("=" * 70)

        print()