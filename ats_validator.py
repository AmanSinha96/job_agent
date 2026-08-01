"""
ats_validator.py
"""

from docx.oxml.ns import qn

MIN_ATS_SCORE = 65


def _structural_issues(document):
    # Keyword coverage alone can't catch a document that many real ATS
    # parsers can't read correctly in the first place — the same class of
    # invisible-in-Word defect as the tab-stop/date-concatenation bug found
    # in resume_base.docx (a docx-internal property with no visual sign in
    # Word, but real ATS-visible impact). This checks for the specific
    # structural patterns known to break plain-text ATS extraction.
    issues = []

    if document.tables:
        issues.append(
            f"{len(document.tables)} table(s) present — many ATS parsers skip table "
            "content entirely or read cells out of order"
        )

    for i, section in enumerate(document.sections):
        cols = section._sectPr.find(qn("w:cols"))
        if cols is not None and cols.get(qn("w:num")) not in (None, "1"):
            issues.append(f"section {i} uses a multi-column layout ({cols.get(qn('w:num'))} columns)")

        header_text = " ".join(p.text for p in section.header.paragraphs).strip()
        footer_text = " ".join(p.text for p in section.footer.paragraphs).strip()
        if header_text or footer_text:
            issues.append(
                f"section {i} has header/footer content ({header_text or footer_text!r}) "
                "— many ATS parsers ignore headers/footers entirely, silently dropping "
                "anything placed there (e.g. name/contact info)"
            )

    if document.inline_shapes:
        issues.append(
            f"{len(document.inline_shapes)} inline image/shape(s) present — not "
            "extractable as text by ATS parsers"
        )

    return issues


def validate_resume(
    resume_text,
    keywords,
    document=None,
):
    """
    Phase 4 compatible validator
    """

    resume_lower = resume_text.lower()

    missing = []

    for kw in keywords:

        if kw.lower() not in resume_lower:

            missing.append(kw)

    matched = (
        len(keywords)
        - len(missing)
    )

    score = 0

    if keywords:

        score = round(
            matched
            / len(keywords)
            * 100,
            2
        )

    structural_issues = _structural_issues(document) if document is not None else []

    return {
        "pass": score >= MIN_ATS_SCORE and not structural_issues,
        "score": score,
        "missing_keywords": missing,
        "structural_issues": structural_issues,
    }