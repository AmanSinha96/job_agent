"""
summary_editor.py

Production Summary Editor

Goals

- Preserve paragraph formatting
- Preserve heading
- Preserve paragraph styles
- Preserve spacing
- Preserve hyperlinks whenever possible
- Only modify summary paragraphs
"""

from docx import Document


SUMMARY_HEADINGS = {

    "professional summary",

    "summary",

    "profile",

}


class SummaryEditor:

    def __init__(

        self,

        document: Document,

    ):

        self.document = document

        self.paragraphs = document.paragraphs

    # --------------------------------------------------

    @staticmethod
    def normalize(

        text,

    ):

        return (

            text

            .replace("\n", " ")

            .strip()

            .lower()

        )

    # --------------------------------------------------

    def is_heading(

        self,

        paragraph,

    ):

        text = self.normalize(

            paragraph.text

        )

        if text in SUMMARY_HEADINGS:

            return True

        style = (

            paragraph.style.name

            .lower()

        )

        if "heading" in style:

            return True

        if (

            paragraph.text.strip()

            == paragraph.text.strip().upper()

            and

            len(paragraph.text.strip()) < 60

        ):

            return True

        return False

    # --------------------------------------------------

    def find_summary_heading(

        self,

    ):

        for index, paragraph in enumerate(

            self.paragraphs

        ):

            if self.normalize(

                paragraph.text

            ) in SUMMARY_HEADINGS:

                return index

        return None

    # --------------------------------------------------

    def find_summary_end(

        self,

        start,

    ):

        for index in range(

            start + 1,

            len(self.paragraphs),

        ):

            if self.is_heading(

                self.paragraphs[index]

            ):

                return index

        return len(

            self.paragraphs

        )

    # --------------------------------------------------

    @staticmethod
    def paragraph_text(

        paragraph,

    ):

        return paragraph.text.strip()

    # --------------------------------------------------

    @staticmethod
    def clear_runs(

        paragraph,

    ):

        """
        Remove only text runs.

        Paragraph formatting
        remains unchanged.
        """

        for run in list(

            paragraph.runs

        ):

            run._element.getparent().remove(

                run._element

            )

    # --------------------------------------------------

    @staticmethod
    def copy_run_style(

        source,

        target,

    ):

        target.bold = source.bold

        target.italic = source.italic

        target.underline = source.underline

        target.font.name = source.font.name

        target.font.size = source.font.size

        target.font.color.rgb = (

            source.font.color.rgb

        )

        target.font.highlight_color = (

            source.font.highlight_color

        )

            # --------------------------------------------------

    def write_text_preserving_style(

        self,

        paragraph,

        text,

    ):
        """
        Replace paragraph text while preserving
        the formatting of the first run.
        """

        if paragraph.runs:

            template = paragraph.runs[0]

            self.clear_runs(

                paragraph

            )

            new_run = paragraph.add_run(

                text

            )

            self.copy_run_style(

                template,

                new_run,

            )

        else:

            paragraph.add_run(

                text

            )

    # --------------------------------------------------

    def replace_summary(

        self,

        new_summary,

    ):
        """
        Replace only the summary section.

        Existing paragraph formatting is preserved.
        """

        heading = self.find_summary_heading()

        if heading is None:

            return False

        end = self.find_summary_end(

            heading

        )

        summary_paragraphs = self.paragraphs[

            heading + 1:end

        ]

        if not summary_paragraphs:

            return False

        new_lines = [

            line.strip()

            for line in new_summary.splitlines()

            if line.strip()

        ]

        if not new_lines:

            return False

        # ------------------------------------------
        # Update existing paragraphs
        # ------------------------------------------

        for index, paragraph in enumerate(

            summary_paragraphs

        ):

            if index < len(new_lines):

                self.write_text_preserving_style(

                    paragraph,

                    new_lines[index],

                )

            else:

                self.write_text_preserving_style(

                    paragraph,

                    "",

                )

        # ------------------------------------------
        # Add additional paragraphs if needed
        # ------------------------------------------

        if len(new_lines) > len(summary_paragraphs):

            anchor = summary_paragraphs[-1]

            for line in new_lines[

                len(summary_paragraphs):

            ]:

                new_para = anchor.insert_paragraph_before()

                new_para.style = anchor.style

                new_run = new_para.add_run(

                    line

                )

                if anchor.runs:

                    self.copy_run_style(

                        anchor.runs[0],

                        new_run,

                    )

                anchor = new_para

        return True

    # --------------------------------------------------

    def update(

        self,

        summary,

    ):
        """
        Public API.
        """

        return self.replace_summary(

            summary

        )


# ======================================================
# Standalone Test
# ======================================================

if __name__ == "__main__":

    from docx import Document

    doc = Document(

        "master_resume.docx"

    )

    editor = SummaryEditor(

        doc

    )

    editor.update(

        "Results-driven Data Engineer with expertise in Python, SQL, Snowflake, Databricks, Airflow, AWS, and Generative AI. Proven experience building scalable ETL pipelines, cloud-native data platforms, and AI-powered automation solutions while improving data quality, operational efficiency, and business insights."

    )

    doc.save(

        "summary_test.docx"

    )

    print(

        "Summary updated successfully."

    )